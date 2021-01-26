from django.shortcuts import render,redirect
from django.contrib.auth import login, authenticate, logout as django_logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from .models import *
from core.forms import *
import json,datetime

# Funcion principal de inicio para la pagina
def inicio(request):
    contexto={
        'titulo': "Inicio",
    }
    return render(request,"inicio.html",contexto)

# Vista para que el alumno pueda verificar su informacion
def login_alumno(request):
    if request.POST and request.is_ajax:
        print("Matricula: "+str(request.POST.get('matricula')))
        # Condicion que permite validar si realmente existen datos de expediente para mostrar al alumno
        if Expediente.objects.filter(matricula=request.POST.get('matricula')).exists():
            return HttpResponse('redirect')
        else:
            return JsonResponse({'type':'info','message':'La matricula ingresada no tiene un expediente activo'})
    contexto={
        'titulo': "Login"
    }
    return render(request,"alumno/login_alumno.html",contexto)

# Validacion de entrada para usuarios de tipo docentes
def login_docente(request):
    logout(request)
    if request.POST:
        # Comprueba si el usuario que se esta logeando existe en la base de datos
        user = request.POST['username']
        password = request.POST['password']
        user = authenticate(username=user, password=password)
        if user is not None:   
            # Realiza el login del usuario
            login(request, user)
            """ Condicion en la cual se verifica si al logearse el usuario existe un perfil ya creado el cual
                sirve para realizar el registro de informacion dando como acceso para realizar las operaciones
                necesarias en el sistema 
            """
            if Perfil.objects.filter(user=request.user.id).exists():
                data=Perfil.objects.get(user=request.user.id)
                if data.estado:
                    return redirect("inicio_docente")
                else: 
                    return redirect("dashbord")
            else:
                usuario=User.objects.get(id=request.user.id)
                model=Perfil(user=usuario, nombre="",apellido_paterno="", apellido_materno="",email_institucional="")
                model.save() 
                return redirect("dashbord")
        else:
            messages.info(request,"Verifica el usuario y contraseña")
    else:
        return render(request,'docente/login_docente.html')

    contexto={
        'title':'Inicio de sesión',
    } 
    return render(request, 'docente/login_docente.html', contexto)

# Vista principal que funciona para deslogear al usuario del sistema
def logout(request):
    django_logout(request)
    return redirect("/")

@login_required(login_url='login_docente')
# Vista que tiene como funcion dar de alta la informacion de los perfiles de los usuarios
def dashbord_docente(request):
    data=Perfil.objects.get(user=request.user.id)
    """ Se genera una instancia sobre la informacion del perfil previamente creado, la cual permite pasar el id del usuario
        al form deseado """
    form=AltaPerfilForm(instance=data)
    if request.POST:
        # Parametros necesarios para el registro de informacion, archivos e identificador de usuario
        form=AltaPerfilForm(request.POST,request.FILES,instance=data)
        print(request.POST)
        """ Si el form contiene informacion correcta se guarda el registro y se cambia el estado del perfil de usuario 
            la cual permitira realizar validaciones al momento de realizar un login """
        if form.is_valid():
            form.save()
            data.estado=True
            data.save()
            print("Registro guardado correctamente")
            messages.success(request,'Perfil creado correctamente',"alert-success")
            return redirect('inicio_docente')
        else:
            print(form.errors)
            messages.warning(request,'Error durante el guardado de datos',"alert-danger")
            return redirect('dashbord')
    contexto={
        'titulo': "Dashboard",
        'usuario': Perfil.objects.get(user=request.user.id),
        'form':form
    }
    return render(request,"docente/dashbord.html",contexto)

@login_required(login_url='login_docente')
# Vista del incio para los usuario con permisos de tipo docente
def inicio_docente(request):
    contexto={
        'titulo': "Inicio",
        'usuario': Perfil.objects.get(user=request.user.id)
    }
    return render(request,"docente/inicio.html",contexto)

@login_required(login_url='login_docente')
# Funcion utilizada para poder editar un perfil del usuario ya registrado
def editar_perfil(request):
    data=Perfil.objects.get(user=request.user.id)
    form=EditarPerfilForm(instance=data)
    url_image=""
    url_name=""
    # Condicion que permite realizar el retorno visual para el usuario del perfil evitando errores
    if data.foto:
        url_image=data.foto.url
        url_name=data.foto
    if request.POST:
        form=EditarPerfilForm(request.POST,request.FILES,instance=data)
        if form.is_valid():
            form.save()
            print("Registro guardado correctamente")
            messages.success(request,'Perfil actualizado correctamente',"alert-success")
            return redirect('editar_perfil')
        else:
            messages.warning(request,'Error durante el guardado de datos',"alert-danger")
    contexto={
        'form':form,
        'titulo': "Editar Perfil",
        'usuario': Perfil.objects.get(user=request.user.id),
        'url_name': url_name,
        'url_image': url_image
    }
    return render(request,"docente/perfil.html",contexto)

@login_required(login_url='login_docente')
# Funcion que permite el retorno de expedientes de creditos filtrados por el usuario logeado y ordenados por el numero de matricula
def expediente(request):  
    expedientes=Expediente.objects.filter(perfil__user=request.user.id,estado=True).order_by("matricula")
    print(expedientes)
    contexto={
        'titulo': "Expedientes",
        'usuario': Perfil.objects.get(user=request.user.id),
        'expedientes':expedientes,
        'total_expedientes': expedientes.count()
    }
    return render(request,"docente/expediente.html",contexto)

@login_required(login_url='login_docente')
# Funcion que permite el retorno de expedientes de servicio social filtrados por el usuario logeado y ordenados por el numero de matricula
def expedienteServicio(request):  
    expedientes=Expediente_Servicio.objects.filter(perfil__user=request.user.id,estado=True).order_by("matricula")
    print(expedientes)
    contexto={
        'titulo': "Expedientes Servicio Social",
        'usuario': Perfil.objects.get(user=request.user.id),
        'expedientes':expedientes,
        'total_expedientes': expedientes.count()
    }
    return render(request,"servicio/expediente.html",contexto)

@login_required(login_url='login_docente')
# Funcion que recibe un parametro id la cual permite filtrar los creditos dentro del expediente registrados  
def detalle_expediente(request,id=None):
    creditos=Credito.objects.filter(user=request.user.id, matricula=id)
    modulos=Perfil.objects.get(user=request.user.id)
    contexto={
        'titulo': "Detalle Expediente",
        'modulos':modulos.modulo,
        'usuario': Perfil.objects.get(user=request.user.id),
        'creditos': creditos,
    }
    return render(request,"docente/detalle_expediente.html",contexto)

# Funcion que recibe un parametro id la cual permite filtrar los documentos del servicio dentro del expediente registrados  
@login_required(login_url='login_docente')
def detalle_expediente_servicio(request,id=None):
    documentos=Documento_Servicio.objects.filter(user=request.user.id, matricula=id)
    modulos=Perfil.objects.get(user=request.user.id)
    contexto={
        'titulo': "Detalle Expediente Servicio Social",
        'modulos':modulos.modulo,
        'usuario': Perfil.objects.get(user=request.user.id),
        'documentos': documentos,
    }
    return render(request,"servicio/detalle_expediente.html",contexto)

@login_required(login_url='login_docente')
# Funcion especial que permite a los usuarios agregar nuevos expedientes de servicio social por metodo ajax 
def Ajax_AltaExpedienteServicio(request):
    user=User.objects.get(id=request.user.id)
    perfil=Perfil.objects.get(user=user)
    form=Expediente_Servicio()
    mensaje=""
    tipo=""
    identify=0
    conteo=0
    if request.is_ajax and request.POST:
        # Validacion para verificar si existe el expediente en el sistema
        if Expediente_Servicio.objects.filter(matricula=request.POST['matricula'], perfil=perfil).exists():
            print("Matricula ya existe no se puede agregar")
            mensaje="Expediente existente no se puede agregar"
            tipo="info"
            identify=0
        else:
            # Si no existe un expediente se registran los datos enviados retornando la matricula para mostrar al usuario
            if request.POST['matricula'] != "" and request.POST['nombre'] != "" and request.POST['apellidoP'] != "" and request.POST['apellidoM'] != "":
                print("Agregando matricula")
                form.matricula=request.POST['matricula']
                form.nombre=request.POST['nombre']
                form.apellido_paterno=request.POST['apellidoP']
                form.apellido_materno=request.POST['apellidoM']
                form.perfil=perfil
                form.save()
                print("Agregado correctamente")
                mensaje="Expediente registrado correctamente"   
                tipo="success"
                conteo=Expediente_Servicio.objects.filter(perfil__user=request.user.id).count()
                print("Este es el conteo: "+str(conteo))
                identify=Expediente_Servicio.objects.last()
                identify=identify.id
            else:
                mensaje="Ha ocurrido un error verifique la informacion ingresada"
                tipo="error"
    return JsonResponse({'mensaje':mensaje,'tipo':tipo,'conteo':conteo,'id':identify})


@login_required(login_url='login_docente')
# Funcion especial que permite a los usuarios agregar nuevos expedientes de creditos por metodo ajax 
def Ajax_AltaExpediente(request):
    user=User.objects.get(id=request.user.id)
    perfil=Perfil.objects.get(user=user)
    form=Expediente()
    mensaje=""
    tipo=""
    identify=0
    conteo=0
    if request.is_ajax and request.POST:
        # Validacion para verificar si existe el expediente en el sistema
        if Expediente.objects.filter(matricula=request.POST['matricula'], perfil=perfil).exists():
            print("Matricula ya existe no se puede agregar")
            mensaje="Expediente existente no se puede agregar"
            tipo="info"
            identify=0
        else:
            # Si no existe un expediente se registran los datos enviados retornando la matricula para mostrar al usuario
            if request.POST['matricula'] != "" and request.POST['nombre'] != "" and request.POST['apellidoP'] != "" and request.POST['apellidoM'] != "":
                print("Agregando matricula")
                form.matricula=request.POST['matricula']
                form.nombre=request.POST['nombre']
                form.apellido_paterno=request.POST['apellidoP']
                form.apellido_materno=request.POST['apellidoM']
                form.perfil=perfil
                form.save()
                print("Agregado correctamente")
                mensaje="Expediente registrado correctamente"   
                tipo="success"
                conteo=Expediente.objects.filter(perfil__user=request.user.id).count()
                print("Este es el conteo: "+str(conteo))
                identify=Expediente.objects.last()
                identify=identify.id
            else:
                mensaje="Ha ocurrido un error verifique la informacion ingresada"
                tipo="error"
    return JsonResponse({'mensaje':mensaje,'tipo':tipo,'conteo':conteo,'id':identify})

@login_required(login_url='login_docente')
# Funcion que permite registrar los documentos en los expedientes por metodo ajax y dropzone
def Ajax_AltaDocumentoExpediente(request):
    user=User.objects.get(id=request.user.id)
    if request.POST and request.is_ajax:
        form=AltaCredito(request.POST, request.FILES or None)
        print(request.POST)
        if form.is_valid():
            # Se obtienen los valores del dropzone los cuales son el rubro, nombre del arcivo y el docuento a registrar
            print("Este es el documento"+str(request.FILES.get('directorio')))
            objeto=form.save(commit=False)
            objeto.nombre=request.FILES.get('directorio')
            objeto.rubro=request.POST.get('rubro')
            objeto.user=user
            form.save()
            modelo=Credito.objects.filter(matricula=request.POST.get('matricula')).last()
            print("Registro guardado correctamente")
            # Al momento de registrar el documento se retornan datos del archivo los cuales se mostraran en una tabla para los usuarios
            return JsonResponse(data={'type':'success','message':'Documento agregado correctamente','url':str(modelo.directorio.url),'nombre':str(modelo.nombre),'rubro':str(modelo.rubro)})
        else:
            print(form.errors)
    return HttpResponse("OK")

@login_required(login_url='login_docente')
# Funcion que permite registrar los documentos en los expedientes de servicio por metodo ajax y dropzone
def Ajax_AltaDocumentoExpedienteServicio(request):
    user=User.objects.get(id=request.user.id)
    if request.POST and request.is_ajax:
        form=AltaDocumento(request.POST, request.FILES or None)
        print(request.POST)
        if form.is_valid():
            print("Este es el documento"+str(request.FILES.get('directorio')))
            objeto=form.save(commit=False)
            objeto.nombre=request.FILES.get('directorio')
            objeto.user=user
            form.save()
            modelo=Documento_Servicio.objects.filter(matricula=request.POST.get('matricula')).last()
            print("Registro guardado correctamente")
            return JsonResponse(data={'type':'success','message':'Documento agregado correctamente','url':str(modelo.directorio.url),'nombre':str(modelo.nombre)})
        else:
            print(form.errors)
    return HttpResponse("OK")

def expedienteAlumno(request,matricula=None):    
    print(str(matricula))
    data=Credito.objects.filter(matricula__matricula=matricula)
    documentos=Documento_Servicio.objects.filter(matricula__matricula=matricula)
    print("Estos son los creditos: "+str(data))
    print("Estos son los documentos: "+str(documentos))

    contexto={
        'title':'Busqueda Expediente',
        'exp':data,
        'documentos':documentos,
    }
    return render(request,"alumno/detalle_expediente.html",contexto)

@login_required(login_url='login_docente')
# Vista especial para mostrar el historial definida por el usuario
def historial(request):
    option=Modulo.objects.all()
    contexto={
        'option':option,
        'usuario': Perfil.objects.get(user=request.user.id),
        'title':'Historial',
        'titulo': "Historial",
    }
    return render(request,'docente/historial.html',contexto)

@login_required(login_url='login_docente')
# Vista para realizar el alta de los modulos/rubros nuevos en el sistema asignandoles un nombre y un valor
def modulo(request):
    if request.is_ajax and request.POST:
        if request.POST['tipo'] == "add":
            modelo=Modulo()
            modelo.tipo=request.POST['modulo']
            modelo.maximo=request.POST['valor']
            modelo.save()
            info=Modulo.objects.get(tipo=request.POST['modulo'])
            return JsonResponse({'id':info.id})
        elif request.POST['tipo'] == "update":
            Modulo.objects.filter(id=request.POST['id']).delete()
    contexto={
        'title':'Modulo',
        'titulo':'Modulo',
        'usuario': Perfil.objects.get(user=request.user.id),
        'data':Modulo.objects.all(),
    }
    return render(request,'docente/modulo.html',contexto)

@login_required(login_url='login_docente')
# Vista especial que permite dar de alta los nuevos creditos a redactar por los usuarios
def alta_credito(request):
    plantillas=None
    modulos=Perfil.objects.get(user=request.user.id)
    if Plantilla.objects.filter(user=request.user.id).exists():
        plantillas=Plantilla.objects.filter(user=request.user.id)
    contexto={
        'title':'Redactar credito',
        'titulo':'Redactar credito',
        'modulos':modulos.modulo,
        'usuario': Perfil.objects.get(user=request.user.id),
        'plantillas':plantillas,
    }
    return render(request,'docente/alta_credito.html',contexto)

@login_required(login_url='login_docente')
# Metodo ajax de consulta de historial con filtros especiales
def Ajax_ConsultaHistorial(request):
    if request.is_ajax and request.POST:
        tiempo=datetime.datetime.strptime(request.POST.get('date_end'), "%Y-%m-%d").date()
        dia=tiempo+datetime.timedelta(days=1)
        print(dia)
        # Se convierte en json lo que recibimos de la vista ya que permite un mejor manejo de informacion
        array=json.loads(request.POST.get('module'))
        # Se filtan los expedientes creados con los filtros asignados de fecha y modulos
        data=Expediente.objects.filter(perfil__modulo__tipo__in=array).filter(created_at__range=[str(request.POST.get('date_start')),str(dia)])
        
        print("Esta es la data "+str(data))
    return HttpResponse("ok")

# Librerias necesarias para poder realizar el proceso de conversion de la plantilla word a pdf y a la modificacion de informacion
from docx import Document
#import locale
#from django.core.files import File as DjangoFile
from os import remove
import shutil
#from docx2pdf import convert
#import pythoncom
from django.db.models import Sum
from django.conf import settings
from django.http import HttpResponse
import os.path as path 

@login_required(login_url='login_docente')
# Ajax que se ejecuta cuando se redacta un nuevo credito por los usuarios
def Ajax_RedactarCredito(request):
    conteo={}
    tipo=""
    mensaje=""
    doc=""
    if request.is_ajax and request.POST:
        # Inicializacion de funcion especial para realizar la modificacion de la plantilla
        #pythoncom.CoInitialize()
        #nombre_doc=request.POST.get("matricula")+"_"+request.POST.get("nombre")
        # Si existe una plantilla dentro del sistema se tomara al usuario para que realice su eleccion de una de ellas
        if Plantilla.objects.filter(id=request.POST.get("select_docplantilla")).exists():
            data=Plantilla.objects.get(id=request.POST.get("select_docplantilla"))
            doc=Document(data.archivo.path)
            print("Este es el doc "+str(doc))
            print("Esta entrando en la plantilla predeterminada")
        # En este caso se toma la plantilla cargada por el usuario
        else:
            print("Esta entrando en la plantilla definida")
            print("Este es el doc "+str(doc))
            doc=Document(request.FILES.get("documento"))
        # Se obtiene la suma de los creditos redactados por el modulo seleccionado al momento de redactar
        if Historial_Redaccion.objects.filter(matricula=request.POST.get("matricula"), modulo=request.POST.get("modulo")).exists():
            conteo=Historial_Redaccion.objects.filter(matricula=request.POST.get("matricula"), modulo=request.POST.get("modulo")).aggregate(Sum('valor'))
        else:
            conteo={"valor__sum":0}
        data=Modulo.objects.get(tipo=str(request.POST.get("modulo")))
        print("Este es el conteo: "+str(conteo)+" y el limite: "+str(data.maximo))
        # Realiza la operacion de diferencia entre el valor maximo de registro de creditos y el valor que se tiene ya dados de alta
        diferencia=int(data.maximo)-int(conteo.get("valor__sum"))
        print("--------- Conteo Historial ---------")
        print("Este es el valor de los creditos asignados: "+str(conteo))
        print("------------------------------------")
        """ Se valida si realmente la diferencia de creditos dados de alta aun no se ha excedido, en caso de que no
            se realizara el proceso de modificacion de documento """
        if diferencia >= int(request.POST.get("valor")):
            today = datetime.datetime.now()
            print("Esta es la fecha: "+str(get_month(today)))
            #locale.setlocale(locale.LC_ALL, ("es_ES", "UTF-8"))
            data=Perfil.objects.get(user=request.user)
            inline=""
            # Se crea un diccionario con sus respectivas claves identicas a la plantilla con su propio valor necesario
            dict={"ENCARGADO":str(data.nombre+" "+data.apellido_paterno+" "+data.apellido_materno),"ALUMNO":request.POST.get("nombre"),"MATRICULA":request.POST.get("matricula"),
            "CARRERA":request.POST.get("carrera"),"RUBRO":request.POST.get("modulo"),"PERIODO":request.POST.get("periodo"),
            "VALOR":request.POST.get("valor"),"DIA":today.strftime("%d"),"MES":str(get_month(today)),"AÑO":today.strftime("%Y"),"JEFA":request.POST.get("jefa")}
            # Recorre las palabras clalves en cada linea de la plantilla
            for values in dict.keys():
                for p in doc.paragraphs:
                    if values in p.text:
                        inline = p.runs
                    # Ciclo que agrega las palabras con cada una de las palabras clave
                    for i in range(len(inline)):
                        if values in inline[i].text:
                            text = inline[i].text.replace(values, dict[values])
                            inline[i].text = text
            # Llama a la funcion a salvar el documento con su respectiva extension
            doc.save("documento.docx")
            print("Se ha guardado la plantilla")
            if path.exists("documento.docx"):
                print("Existe el documento")
                doc.save("documento.docx")
                shutil.move("documento.docx","./"+settings.MEDIA_URL)
                # Cambia el documento de word a pdf
                #convert("documento.docx")
                #convert("documento.docx","documento.pdf")
                #webbrowser.open_new("documento.pdf")
                
            else:
                print("No existe el documento")
            
            # Se remueve el documento word y se pasa el pdf al escritorio del usuario
            #remove(nombre_doc+".docx")
            #os.rename(nombre_doc+".docx",nombre_doc+".docx")
            #desktop = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop') 
            #print("La ruta es la siguiente: "+str(desktop))
            # En caso de que exista el documento lo que se hace es remover el documento antiguo
            #if os.path.isfile(desktop+"/"+nombre_doc+".docx"):
            #    print("Esta entrando a la condicion donde existe el archivo en el escritorio")
            #    shutil.copy(nombre_doc+".docx",desktop)
            #    remove(nombre_doc+".docx")
            #else:
            #    print("Esta entrando a la condicion donde no existe el archivo en el escritorio")
            #    shutil.move(nombre_doc+".docx",desktop)
            # Guardado de historial del credito redactado 
            Historial_Redaccion.objects.create(matricula=request.POST.get('matricula'),nombre=request.POST.get("nombre"),carrera=request.POST.get("carrera"),modulo=request.POST.get("modulo"),valor=request.POST.get("valor"),jefa=request.POST.get('jefa'))
            tipo="success"
            mensaje="realizado"

        #file_obj1 = DjangoFile(open("Plantilla.docx", mode='rb'), name="Plantilla.docx")
        #Plantilla.objects.create(archivo=file_obj1,user=request.user)
        else:
            print("Esta entrando en el else porque sobrepaso los limites")
            tipo="info"
            mensaje="denegado"
        print("Este es el mensaje: "+str(mensaje)+".")
    return JsonResponse({'type':tipo,'result':mensaje})

def get_month(date):
    months = ("Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre")
    month = months[date.month - 1]
    return month

@login_required(login_url='login_docente')
# Funcion ajax que permite guardar la informacion redactada en los creditos para facilitar el registro de un nuevo credito
def Ajax_GuardarHistorialCredito(request):
    if request.is_ajax and request.GET:
        if Historial_Redaccion.objects.filter(matricula=request.GET.get('matricula')).exists():
            data=Historial_Redaccion.objects.filter(matricula=request.GET.get('matricula')).last()
            return JsonResponse({'nombre':str(data.nombre),'carrera':str(data.carrera),'jefa':str(data.jefa)})
        else:
            return HttpResponse("info")
    return HttpResponse("ok")

@login_required(login_url='login_docente')
# Ajax que permite realizar el guardado de las plantillas predeterminadas por los usuarios
def Ajax_GuardarPlantilla(request):
    if request.POST and request.is_ajax:
        print("Entra al metodo para guardar la plantilla")
        usuario=User.objects.get(id=request.user.id)
        Plantilla.objects.create(archivo=request.FILES.get("plantilla"),user=usuario)
        print("Se ha guardado correctamente")
    return HttpResponse("OK")